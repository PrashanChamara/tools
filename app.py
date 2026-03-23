import io
import os
import json
import shutil
import tempfile
import numpy as np
import pandas as pd
import cv2
from flask import Flask, render_template, request, send_file, redirect, url_for, jsonify, session
from docx import Document
from rembg import remove
from PIL import Image, ImageEnhance
from PyPDF2 import PdfReader, PdfWriter
from pdf2image import convert_from_bytes

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'tuxtalk-dev-secret-change-in-prod')

# ── Image Upscaling Setup ─────────────────────────────────
UPSCALE_MODEL_PATH = "models/EDSR_x4.pb"

def upscale_image_with_model(image_path, scale=4):
    sr = cv2.dnn_superres.DnnSuperResImpl_create()
    sr.readModel(UPSCALE_MODEL_PATH)
    sr.setModel("edsr", 4)
    image = cv2.imread(image_path)
    result = sr.upsample(image)
    if scale == 2:
        h, w = result.shape[:2]
        result = cv2.resize(result, (w // 2, h // 2), interpolation=cv2.INTER_LANCZOS4)
    return result

# ── Temp session dir helpers ──────────────────────────────
def get_session_dir():
    sid = session.get('sid')
    if not sid:
        import uuid
        sid = str(uuid.uuid4())
        session['sid'] = sid
    d = os.path.join(tempfile.gettempdir(), f'tuxtalk_{sid}')
    os.makedirs(d, exist_ok=True)
    return d

def cleanup_session_dir():
    d = session.get('_tmpdir')
    if d and os.path.isdir(d):
        shutil.rmtree(d, ignore_errors=True)

# ═══════════════════════════════════════════════════════════
# EXISTING ROUTES (unchanged)
# ═══════════════════════════════════════════════════════════

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/speech_to_text')
def speech_to_text():
    return render_template('speech_to_text.html')

@app.route('/download', methods=['POST'])
def download():
    text = request.form.get('text', '')
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return send_file(file_stream, as_attachment=True,
                     download_name="transcription.docx",
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/remove_background')
def remove_background():
    return render_template('remove_background.html')

@app.route('/remove_bg', methods=['POST'])
def remove_bg():
    file = request.files.get('image')
    if not file:
        return "No file uploaded", 400
    input_image = Image.open(file).convert("RGBA")
    output_image = remove(input_image)
    output_stream = io.BytesIO()
    output_image.save(output_stream, format="PNG")
    output_stream.seek(0)
    return send_file(output_stream, as_attachment=True, download_name="no_bg.png", mimetype="image/png")

@app.route('/sketch')
def sketch():
    return render_template('sketch.html')

@app.route('/sketch_image', methods=['POST'])
def sketch_image():
    file = request.files.get('image')
    if not file:
        return "No file uploaded", 400
    try:
        intensity = int(request.form.get('intensity', 21))
        # Ensure odd kernel size
        if intensity % 2 == 0:
            intensity += 1
        intensity = max(3, min(51, intensity))
    except (ValueError, TypeError):
        intensity = 21

    file_bytes = np.asarray(bytearray(file.read()), dtype=np.uint8)
    img = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
    if img is None:
        return "Invalid image", 400
    gray_img     = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    inverted_img = cv2.bitwise_not(gray_img)
    blurred      = cv2.GaussianBlur(inverted_img, (intensity, intensity), sigmaX=0, sigmaY=0)
    inverted_blur= cv2.bitwise_not(blurred)
    sketch_img   = cv2.divide(gray_img, inverted_blur, scale=256.0)
    success, buffer = cv2.imencode('.png', sketch_img)
    if not success:
        return "Image processing failed", 500
    return send_file(io.BytesIO(buffer), as_attachment=True, download_name="sketch.png", mimetype="image/png")

# ── PDF Tool ──────────────────────────────────────────────
@app.route('/pdf_tool', methods=['GET', 'POST'])
def pdf_tool():
    if request.method == 'GET':
        return render_template('pdf_tool.html')

    files     = request.files.getlist('pdf_files')
    operation = request.form.get('operation')
    if not files:
        return "No PDF files uploaded", 400

    pdf_data = [f.read() for f in files if f.filename]
    if not pdf_data:
        return "No valid PDF files", 400

    rotate_option = request.form.get('rotate_option')

    if operation == 'merge':
        writer = PdfWriter()
        for data in pdf_data:
            reader = PdfReader(io.BytesIO(data))
            for page in reader.pages:
                writer.add_page(page)
        output_stream = io.BytesIO()
        writer.write(output_stream)
        output_stream.seek(0)
        return send_file(output_stream, as_attachment=True, download_name="merged.pdf", mimetype="application/pdf")

    elif operation == 'split':
        reader = PdfReader(io.BytesIO(pdf_data[0]))
        writer = PdfWriter()
        if reader.pages:
            writer.add_page(reader.pages[0])
        out = io.BytesIO()
        writer.write(out); out.seek(0)
        return send_file(out, as_attachment=True, download_name="split_first_page.pdf", mimetype="application/pdf")

    elif operation == 'pdf_to_jpg':
        images = convert_from_bytes(pdf_data[0])
        if not images:
            return "Conversion failed", 500
        img_io = io.BytesIO()
        images[0].save(img_io, 'JPEG')
        img_io.seek(0)
        return send_file(img_io, as_attachment=True, download_name="converted.jpg", mimetype="image/jpeg")

    elif operation == 'pdf_to_word':
        try:
            reader = PdfReader(io.BytesIO(pdf_data[0]))
            text_parts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
            if not text_parts:
                return "This PDF contains no extractable text. Scanned PDFs are not supported.", 422
            doc = Document()
            for part in text_parts:
                for line in part.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
                doc.add_page_break()
            out = io.BytesIO()
            doc.save(out); out.seek(0)
            return send_file(out, as_attachment=True, download_name="extracted.docx",
                             mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except Exception as e:
            return f"Conversion failed: {str(e)}", 500

    elif operation == 'pdf_to_excel':
        try:
            reader = PdfReader(io.BytesIO(pdf_data[0]))
            rows = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    for line in t.split('\n'):
                        if line.strip():
                            cols = line.split()
                            rows.append(cols)
            if not rows:
                return "This PDF contains no extractable text. Scanned PDFs are not supported.", 422
            max_cols = max(len(r) for r in rows)
            for r in rows:
                while len(r) < max_cols:
                    r.append('')
            df = pd.DataFrame(rows)
            out = io.BytesIO()
            df.to_excel(out, index=False, header=False)
            out.seek(0)
            return send_file(out, as_attachment=True, download_name="extracted.xlsx",
                             mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        except Exception as e:
            return f"Conversion failed: {str(e)}", 500

    elif operation == 'rotate':
        try:
            angle = int(rotate_option or 90)
        except ValueError:
            angle = 90
        reader = PdfReader(io.BytesIO(pdf_data[0]))
        writer = PdfWriter()
        for page in reader.pages:
            page.rotate(angle)
            writer.add_page(page)
        out = io.BytesIO()
        writer.write(out); out.seek(0)
        return send_file(out, as_attachment=True, download_name="rotated.pdf", mimetype="application/pdf")

    return "Invalid operation", 400

# ── Image Upscaling ───────────────────────────────────────
@app.route('/image_upscaling')
def image_upscaling():
    return render_template('image_upscaling.html')

@app.route('/process_upscaling', methods=['POST'])
def process_upscaling():
    file = request.files.get('image')
    if not file:
        return "No file uploaded", 400

    try:
        scale = int(request.form.get('scale', 4))
        if scale not in (2, 4):
            scale = 4
    except (ValueError, TypeError):
        scale = 4

    temp_filename = os.path.join(tempfile.gettempdir(), 'tuxtalk_upscale_tmp.png')
    file.save(temp_filename)

    img_check = cv2.imread(temp_filename)
    if img_check is None:
        os.remove(temp_filename)
        return "Invalid image file.", 400

    h, w = img_check.shape[:2]
    if h * w > 1_000_000:
        os.remove(temp_filename)
        return (f"Image is too large ({w}×{h}). Please use an image smaller than 1 megapixel.", 413)

    if not os.path.exists(UPSCALE_MODEL_PATH):
        os.remove(temp_filename)
        return "Upscaling model not found on the server.", 500

    try:
        result = upscale_image_with_model(temp_filename, scale=scale)
        os.remove(temp_filename)
        success, buffer = cv2.imencode('.png', result)
        if not success:
            return "Image processing failed", 500
        return send_file(io.BytesIO(buffer), as_attachment=True,
                         download_name=f"upscaled_{scale}x.png", mimetype="image/png")
    except Exception as e:
        if os.path.exists(temp_filename):
            os.remove(temp_filename)
        return f"Upscaling failed: {str(e)}", 500

# ═══════════════════════════════════════════════════════════
# EXCEL MERGE — Enhanced (multi-file)
# ═══════════════════════════════════════════════════════════

@app.route('/excel_tool', methods=['GET'])
def excel_tool():
    return render_template('excel_tool.html')

@app.route('/upload_excel_files', methods=['POST'])
def upload_excel_files():
    """Parse a single uploaded Excel file, return column list + preview."""
    file = request.files.get('file')
    if not file:
        return jsonify({'error': 'No file provided'}), 400
    try:
        df = pd.read_excel(io.BytesIO(file.read()))
        if df.empty:
            return jsonify({'error': f'File "{file.filename}" has no data rows.'}), 400
        preview = df.head(3).fillna('').to_dict(orient='records')
        # Convert any non-serialisable types
        for row in preview:
            for k, v in row.items():
                row[k] = str(v) if not isinstance(v, (str, int, float, bool, type(None))) else v
        return jsonify({
            'columns': list(df.columns),
            'rows':    len(df),
            'preview': preview
        })
    except Exception as e:
        return jsonify({'error': f'Could not read "{file.filename}": {str(e)}'}), 400

@app.route('/process_excel_multi', methods=['POST'])
def process_excel_multi():
    """Merge up to 5 Excel files based on a config payload."""
    config_raw = request.form.get('config')
    if not config_raw:
        return "Missing merge configuration", 400

    try:
        config = json.loads(config_raw)
    except json.JSONDecodeError:
        return "Invalid configuration format", 400

    key_col   = config.get('key_col')
    join_type = config.get('join_type', 'left')
    files_cfg = config.get('files', [])

    if not key_col:
        return "No join key column specified", 400
    if join_type not in ('left', 'inner', 'outer'):
        join_type = 'left'

    # Read all uploaded files
    dfs = []
    letters = []
    for fc in files_cfg:
        idx    = fc.get('index', 0)
        letter = fc.get('letter', str(idx))
        sel    = fc.get('selected_cols', [])
        f      = request.files.get(f'file_{idx}')
        if not f:
            continue
        try:
            df = pd.read_excel(io.BytesIO(f.read()))
        except Exception as e:
            return f"Could not read file {idx}: {str(e)}", 400

        if key_col not in df.columns:
            return f"File {letter.upper()} does not have column '{key_col}'.", 400

        # Keep only selected cols + key
        keep = [key_col] + [c for c in sel if c in df.columns and c != key_col]
        df   = df[keep]
        dfs.append(df)
        letters.append(letter)

    if not dfs:
        return "No files to merge", 400
    if len(dfs) == 1:
        result = dfs[0]
    else:
        result = dfs[0]
        for i, df in enumerate(dfs[1:], start=1):
            suffix_left  = f'_{letters[0]}'
            suffix_right = f'_{letters[i]}'
            result = pd.merge(result, df, on=key_col, how=join_type,
                              suffixes=(suffix_left, suffix_right))

    if result.empty and join_type == 'inner':
        warning = f"No rows matched on column '{key_col}'. Try Left or Outer join."
        out = io.BytesIO()
        result.to_excel(out, index=False)
        out.seek(0)
        response = send_file(out, as_attachment=True, download_name="merged.xlsx",
                             mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        response.headers['X-Warning'] = warning
        return response

    out = io.BytesIO()
    result.to_excel(out, index=False)
    out.seek(0)
    return send_file(out, as_attachment=True, download_name="merged.xlsx",
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# Redirect old route
@app.route('/process_excel', methods=['POST'])
def process_excel():
    return redirect(url_for('excel_tool'))

# ═══════════════════════════════════════════════════════════
# COLOR CORRECTION — Studio
# ═══════════════════════════════════════════════════════════

@app.route('/color_correction')
def color_correction():
    return render_template('color_correction.html')

def _apply_color_correction(img_bgr, sliders, curves_data):
    """Apply the full studio pipeline. Returns corrected BGR image."""
    img = img_bgr.astype(np.float32)

    # 1. Exposure (EV — stored as -200..200 → actual EV -2..+2)
    ev = sliders.get('exposure', 0) / 100.0
    img = img * (2 ** ev)
    img = np.clip(img, 0, 255)

    # 2. Whites / Blacks (simple tone-range clip/push)
    whites = sliders.get('whites', 0) / 100.0
    blacks = sliders.get('blacks', 0) / 100.0
    if whites != 0:
        img = img + (255 - img) * (whites * 0.5 if whites > 0 else whites * 0.3)
    if blacks != 0:
        img = img + img * (blacks * 0.3 if blacks < 0 else blacks * 0.5)
    img = np.clip(img, 0, 255)

    # 3. Highlights / Shadows via LAB
    img_u8 = img.astype(np.uint8)
    lab = cv2.cvtColor(img_u8, cv2.COLOR_BGR2LAB).astype(np.float32)
    L = lab[:,:,0]

    highlights = sliders.get('highlights', 0) / 100.0
    shadows    = sliders.get('shadows',    0) / 100.0

    if highlights != 0:
        mask_hi = np.clip((L - 128) / 128, 0, 1)
        L = L + highlights * 30 * mask_hi
    if shadows != 0:
        mask_lo = np.clip(1 - L / 128, 0, 1)
        L = L + shadows * 30 * mask_lo

    lab[:,:,0] = np.clip(L, 0, 255)
    img = cv2.cvtColor(lab.astype(np.uint8), cv2.COLOR_LAB2BGR).astype(np.float32)

    # 4. Contrast via Pillow
    contrast = sliders.get('contrast', 0) / 100.0
    if contrast != 0:
        pil_img = Image.fromarray(cv2.cvtColor(img.astype(np.uint8), cv2.COLOR_BGR2RGB))
        factor  = 1.0 + contrast * 0.8
        pil_img = ImageEnhance.Contrast(pil_img).enhance(max(0.1, factor))
        img = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR).astype(np.float32)

    # 5. Temperature / Tint (channel shift)
    temp = sliders.get('temperature', 0) / 100.0
    tint = sliders.get('tint',        0) / 100.0
    if temp != 0:
        img[:,:,2] = np.clip(img[:,:,2] + temp * 25, 0, 255)  # R channel
        img[:,:,0] = np.clip(img[:,:,0] - temp * 20, 0, 255)  # B channel
    if tint != 0:
        img[:,:,1] = np.clip(img[:,:,1] + tint * 20, 0, 255)  # G channel

    # 6. Curves (LUT per channel)
    img_u8 = img.astype(np.uint8)
    def apply_lut(channel_img, pts):
        if not pts or len(pts) < 2:
            return channel_img
        pts_sorted = sorted(pts, key=lambda p: p[0])
        xs = np.array([p[0] for p in pts_sorted], dtype=np.float32)
        ys = np.array([p[1] for p in pts_sorted], dtype=np.float32)
        xs = np.clip(xs, 0, 255); ys = np.clip(ys, 0, 255)
        lut_in  = np.arange(256, dtype=np.float32)
        lut_out = np.interp(lut_in, xs, ys).astype(np.uint8)
        return lut_out[channel_img]

    if curves_data:
        cr = curves_data.get('curves_r',   [[0,0],[255,255]])
        cg = curves_data.get('curves_g',   [[0,0],[255,255]])
        cb = curves_data.get('curves_b',   [[0,0],[255,255]])
        crgb = curves_data.get('curves_rgb', [[0,0],[255,255]])

        b_ch, g_ch, r_ch = cv2.split(img_u8)
        r_ch = apply_lut(r_ch, cr)
        g_ch = apply_lut(g_ch, cg)
        b_ch = apply_lut(b_ch, cb)
        img_u8 = cv2.merge([b_ch, g_ch, r_ch])

        # Composite RGB curve
        b_ch, g_ch, r_ch = cv2.split(img_u8)
        r_ch = apply_lut(r_ch, crgb)
        g_ch = apply_lut(g_ch, crgb)
        b_ch = apply_lut(b_ch, crgb)
        img_u8 = cv2.merge([b_ch, g_ch, r_ch])

    img = img_u8.astype(np.float32)

    # 7. Clarity (unsharp mask on mid-tones)
    clarity = sliders.get('clarity', 0) / 100.0
    if clarity > 0:
        blurred = cv2.GaussianBlur(img.astype(np.uint8), (0, 0), sigmaX=3)
        img_u8  = img.astype(np.uint8)
        L_mask  = cv2.cvtColor(img_u8, cv2.COLOR_BGR2GRAY).astype(np.float32) / 255.0
        mid_mask= 1 - np.abs(2 * L_mask - 1)  # peaks at 0.5 gray
        sharpened = cv2.addWeighted(img_u8, 1 + clarity * 0.5, blurred, -clarity * 0.5, 0)
        alpha = (mid_mask * clarity * 0.6)[:, :, np.newaxis]
        img = (1 - alpha) * img + alpha * sharpened.astype(np.float32)

    # 8. Vibrance / Saturation (HSV)
    vibrance   = sliders.get('vibrance',   0) / 100.0
    saturation = sliders.get('saturation', 0) / 100.0
    if vibrance != 0 or saturation != 0:
        hsv = cv2.cvtColor(np.clip(img, 0, 255).astype(np.uint8), cv2.COLOR_BGR2HSV).astype(np.float32)
        s = hsv[:,:,1] / 255.0
        if saturation != 0:
            s = np.clip(s * (1 + saturation), 0, 1)
        if vibrance != 0:
            low_sat_mask = 1 - s  # boost low-saturation pixels more
            s = np.clip(s + vibrance * 0.4 * low_sat_mask, 0, 1)
        hsv[:,:,1] = (s * 255).astype(np.float32)
        img = cv2.cvtColor(hsv.astype(np.uint8), cv2.COLOR_HSV2BGR).astype(np.float32)

    return np.clip(img, 0, 255).astype(np.uint8)

def _parse_correction_request():
    file   = request.files.get('image')
    sliders_raw = request.form.get('sliders', '{}')
    curves_raw  = request.form.get('curves',  '{}')
    sliders = json.loads(sliders_raw)
    curves  = json.loads(curves_raw) if curves_raw else {}
    return file, sliders, curves

@app.route('/preview_color_correction', methods=['POST'])
def preview_color_correction():
    file, sliders, curves = _parse_correction_request()
    if not file:
        return "No image", 400
    file_bytes = np.asarray(bytearray(file.read()), dtype=np.uint8)
    img = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
    if img is None:
        return "Invalid image", 400

    # Downscale for fast preview (max 800px wide)
    h, w = img.shape[:2]
    if w > 800:
        scale = 800 / w
        img = cv2.resize(img, (800, int(h * scale)), interpolation=cv2.INTER_AREA)

    result = _apply_color_correction(img, sliders, curves)
    _, buffer = cv2.imencode('.jpg', result, [cv2.IMWRITE_JPEG_QUALITY, 85])
    return send_file(io.BytesIO(buffer), mimetype='image/jpeg')

@app.route('/process_color_correction_studio', methods=['POST'])
def process_color_correction_studio():
    file, sliders, curves = _parse_correction_request()
    if not file:
        return "No image", 400
    file_bytes = np.asarray(bytearray(file.read()), dtype=np.uint8)
    img = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
    if img is None:
        return "Invalid image", 400

    result  = _apply_color_correction(img, sliders, curves)
    _, buffer = cv2.imencode('.png', result)
    return send_file(io.BytesIO(buffer), as_attachment=True,
                     download_name="color_corrected.png", mimetype="image/png")

# Keep old simple route for backward compat
@app.route('/process_color_correction', methods=['POST'])
def process_color_correction():
    file = request.files.get('image')
    mode = request.form.get('mode')
    if not file:
        return "No file uploaded", 400
    file_bytes = np.asarray(bytearray(file.read()), dtype=np.uint8)
    img = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
    if img is None:
        return "Invalid image", 400

    if mode == 'day':
        lab = cv2.cvtColor(img, cv2.COLOR_BGR2LAB)
        l, a, b = cv2.split(lab)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        cl = clahe.apply(l)
        corrected_img = cv2.cvtColor(cv2.merge((cl, a, b)), cv2.COLOR_LAB2BGR)
    elif mode == 'night':
        gamma = 1.5
        lut = np.array([np.clip(pow(i/255.0, 1.0/gamma)*255.0, 0, 255) for i in range(256)], dtype=np.uint8)
        corrected_img = cv2.LUT(img, lut)
    else:
        return "Invalid mode selected", 400

    _, buffer = cv2.imencode('.png', corrected_img)
    return send_file(io.BytesIO(buffer), as_attachment=True,
                     download_name="color_corrected.png", mimetype="image/png")

if __name__ == '__main__':
    app.run(debug=True)
